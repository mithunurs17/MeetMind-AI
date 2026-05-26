import json
from datetime import datetime
from io import BytesIO

import streamlit as st
from docx import Document

from utils.api_client import api_client

# =====================================================
# PAGE CONFIG
# =====================================================

st.set_page_config(
    page_title="MeetMind AI",
    page_icon="🧠",
    layout="wide",
)

# =====================================================
# CUSTOM CSS
# =====================================================

st.markdown("""
<style>

@import url('https://fonts.googleapis.com/css2?family=Instrument+Serif:ital@0;1&family=Barlow:wght@300;400;500;600&display=swap');

html, body, [class*="css"] {
    font-family: 'Barlow', sans-serif;
}

/* =====================================================
HIDE STREAMLIT
===================================================== */

#MainMenu {
    visibility: hidden;
}

footer {
    visibility: hidden;
}

header {
    visibility: hidden;
}

/* =====================================================
APP BACKGROUND
===================================================== */

.stApp {

    background:
        radial-gradient(circle at top left,
        rgba(99,102,241,0.25),
        transparent 25%),

        radial-gradient(circle at bottom right,
        rgba(236,72,153,0.15),
        transparent 25%),

        #000;

    color: white;
}

/* =====================================================
SIDEBAR
===================================================== */

section[data-testid="stSidebar"] {

    background: rgba(255,255,255,0.03);

    backdrop-filter: blur(20px);

    border-right:
        1px solid rgba(255,255,255,0.08);
}

section[data-testid="stSidebar"] * {
    color: white !important;
}

/* =====================================================
LIQUID GLASS
===================================================== */

.liquid-glass {

    background: rgba(255,255,255,0.03);

    background-blend-mode: luminosity;

    backdrop-filter: blur(12px);

    -webkit-backdrop-filter: blur(12px);

    border: none;

    box-shadow:
        inset 0 1px 1px rgba(255,255,255,0.1),
        0 8px 30px rgba(0,0,0,0.35);

    position: relative;

    overflow: hidden;
}

.liquid-glass::before {

    content: "";

    position: absolute;

    inset: 0;

    border-radius: inherit;

    padding: 1.4px;

    background:
        linear-gradient(
            180deg,
            rgba(255,255,255,0.45) 0%,
            rgba(255,255,255,0.15) 20%,
            rgba(255,255,255,0) 40%,
            rgba(255,255,255,0) 60%,
            rgba(255,255,255,0.15) 80%,
            rgba(255,255,255,0.45) 100%
        );

    -webkit-mask:
        linear-gradient(#fff 0 0) content-box,
        linear-gradient(#fff 0 0);

    -webkit-mask-composite: xor;

    mask-composite: exclude;

    pointer-events: none;
}

/* =====================================================
HERO SECTION
===================================================== */

.hero {

    position: relative;

    overflow: hidden;

    border-radius: 36px;

    padding: 80px 60px;

    min-height: 450px;

    background:
        linear-gradient(
            135deg,
            rgba(255,255,255,0.05),
            rgba(255,255,255,0.02)
        );

    margin-bottom: 30px;
}

.hero::after {

    content: "";

    position: absolute;

    width: 600px;
    height: 600px;

    border-radius: 50%;

    background: rgba(255,255,255,0.08);

    top: -300px;
    right: -150px;

    filter: blur(90px);
}

.hero-title {

    font-family: 'Instrument Serif', serif;

    font-style: italic;

    font-size: 5.5rem;

    line-height: 0.9;

    letter-spacing: -4px;

    color: white;

    max-width: 850px;

    position: relative;

    z-index: 2;
}

.hero-sub {

    margin-top: 30px;

    font-size: 17px;

    line-height: 1.6;

    color: rgba(255,255,255,0.82);

    max-width: 700px;

    position: relative;

    z-index: 2;
}

.glass-btn {

    display: inline-flex;

    align-items: center;

    gap: 10px;

    padding: 14px 28px;

    border-radius: 999px;

    margin-top: 35px;

    font-size: 14px;

    font-weight: 500;

    color: white;

    background: rgba(255,255,255,0.04);

    backdrop-filter: blur(20px);

    box-shadow:
        inset 0 1px 1px rgba(255,255,255,0.15),
        0 8px 20px rgba(0,0,0,0.3);

    width: fit-content;

    position: relative;

    z-index: 2;
}

/* =====================================================
METRIC CARDS
===================================================== */

.metric-card {

    border-radius: 28px;

    padding: 28px;

    min-height: 180px;

    transition: 0.35s ease;
}

.metric-card:hover {

    transform: translateY(-8px);
}

.metric-title {

    font-size: 14px;

    color: rgba(255,255,255,0.72);

    margin-bottom: 18px;
}

.metric-value {

    font-family: 'Instrument Serif', serif;

    font-style: italic;

    font-size: 52px;

    line-height: 1;

    color: white;
}

.metric-desc {

    margin-top: 16px;

    font-size: 13px;

    color: rgba(255,255,255,0.72);
}

/* =====================================================
GLASS CONTENT CARD
===================================================== */

.glass-card {

    border-radius: 28px;

    padding: 30px;

    margin-top: 20px;
}

.glass-card h2,
.glass-card h3 {

    color: white;
}

.glass-card p {

    color: rgba(255,255,255,0.82);

    line-height: 1.7;
}

/* =====================================================
PROGRESS BAR
===================================================== */

.progress-bar {

    width: 100%;

    height: 14px;

    border-radius: 999px;

    background: rgba(255,255,255,0.08);

    overflow: hidden;

    margin-top: 20px;
}

.progress-fill {

    width: 70%;

    height: 100%;

    background:
        linear-gradient(
            90deg,
            #06b6d4,
            #8b5cf6,
            #ec4899
        );

    animation: move 3s infinite;
}

@keyframes move {

    0% {
        transform: translateX(-40%);
    }

    100% {
        transform: translateX(100%);
    }
}

/* =====================================================
BUTTONS
===================================================== */

.stDownloadButton > button {

    width: 100%;

    border: none;

    border-radius: 999px;

    background:
        linear-gradient(
            135deg,
            #7c3aed,
            #4f46e5
        );

    color: white;

    font-weight: 600;

    padding: 0.8rem 1rem;

    margin-top: 15px;
}

</style>
""", unsafe_allow_html=True)

# =====================================================
# API DATA
# =====================================================

health_status = api_client.health_check()

meetings_response = api_client.get_meetings()

meetings = meetings_response if isinstance(
    meetings_response,
    list
) else []

action_items = [
    item
    for meeting in meetings
    for item in meeting.get("action_items", [])
]

decisions = [
    item
    for meeting in meetings
    for item in meeting.get("decisions", [])
]

risks = [
    item
    for meeting in meetings
    for item in meeting.get("risks", [])
]

open_questions = [
    q
    for meeting in meetings
    for q in (meeting.get("open_questions") or [])
]

owners = sorted({
    item.get("owner")
    for item in action_items
    if item.get("owner")
})

# =====================================================
# SIDEBAR
# =====================================================

with st.sidebar:

    st.title("🧠 MeetMind AI")

    st.markdown("---")

    st.subheader("Navigation")

    st.markdown("""
- 📹 Upload Meeting
- 📋 View Minutes
- ✅ Action Items
- 📊 Dashboard
""")

    st.markdown("---")

    st.subheader("System Status")

    if health_status.get("status") == "healthy":
        st.success("Backend Connected")
    else:
        st.error("Backend Offline")

# =====================================================
# HERO SECTION
# =====================================================

st.markdown("""
<div class="hero liquid-glass">

    <div class="hero-title">
        Venture Past <br>
        Our Sky Across <br>
        the Universe
    </div>

    <div class="hero-sub">
        AI-powered meeting intelligence platform that transforms
        transcripts into summaries, action items,
        decisions, risks, and premium analytics in real time.
    </div>

    <div class="glass-btn">
        🚀 Start Your Voyage
    </div>

</div>
""", unsafe_allow_html=True)

# =====================================================
# METRICS
# =====================================================

col1, col2, col3, col4 = st.columns(4)

with col1:

    st.markdown(f"""
    <div class="metric-card liquid-glass">

        <div class="metric-title">
            Total Meetings
        </div>

        <div class="metric-value">
            {len(meetings)}
        </div>

        <div class="metric-desc">
            Meetings analyzed
        </div>

    </div>
    """, unsafe_allow_html=True)

with col2:

    st.markdown(f"""
    <div class="metric-card liquid-glass">

        <div class="metric-title">
            Action Items
        </div>

        <div class="metric-value">
            {len(action_items)}
        </div>

        <div class="metric-desc">
            Tasks extracted
        </div>

    </div>
    """, unsafe_allow_html=True)

with col3:

    st.markdown(f"""
    <div class="metric-card liquid-glass">

        <div class="metric-title">
            Decisions
        </div>

        <div class="metric-value">
            {len(decisions)}
        </div>

        <div class="metric-desc">
            Key decisions captured
        </div>

    </div>
    """, unsafe_allow_html=True)

with col4:

    st.markdown(f"""
    <div class="metric-card liquid-glass">

        <div class="metric-title">
            Risks & Questions
        </div>

        <div class="metric-value">
            {len(risks) + len(open_questions)}
        </div>

        <div class="metric-desc">
            Pending blockers
        </div>

    </div>
    """, unsafe_allow_html=True)

# =====================================================
# ANALYTICS SECTION
# =====================================================

left, right = st.columns([3, 1.5])

with left:

    st.markdown("""
    <div class="glass-card liquid-glass">

        <h2>📈 Meeting Analytics</h2>

        <p>
            Visualize collaboration efficiency and
            meeting insights in real-time using
            AI-powered analysis.
        </p>

        <div class="progress-bar">
            <div class="progress-fill"></div>
        </div>

        <p style="margin-top:20px;">
            Keep uploading meeting transcripts to unlock richer
            analytics and productivity insights.
        </p>

    </div>
    """, unsafe_allow_html=True)

with right:

    st.markdown(f"""
    <div class="glass-card liquid-glass">

        <h3>⚡ Quick Insights</h3>

        <p><b>👥 Owners Engaged:</b> {len(owners)}</p>

        <p><b>🟢 AI Health:</b> {health_status.get('status', 'unknown')}</p>

        <p><b>🕒 Last Sync:</b><br>
        {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}</p>

    </div>
    """, unsafe_allow_html=True)

    # =====================================================
    # DOCX EXPORT
    # =====================================================

    doc = Document()

    doc.add_heading(
        "MeetMind AI - Meeting Report",
        level=1
    )

    for index, meeting in enumerate(meetings, start=1):

        doc.add_heading(
            f"Meeting {index}",
            level=2
        )

        if meeting.get("summary"):

            doc.add_heading(
                "Summary",
                level=3
            )

            doc.add_paragraph(
                str(meeting.get("summary"))
            )

        if meeting.get("action_items"):

            doc.add_heading(
                "Action Items",
                level=3
            )

            for item in meeting["action_items"]:

                task = item.get(
                    "task",
                    "No Task"
                )

                owner = item.get(
                    "owner",
                    "Unknown"
                )

                doc.add_paragraph(
                    f"• {task} (Owner: {owner})"
                )

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    st.download_button(
        label="⬇ Export Meetings DOCX",
        data=buffer,
        file_name="meetmind_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# =====================================================
# GETTING STARTED
# =====================================================

st.markdown("""
<div class="glass-card liquid-glass">

    <h2>🚀 Getting Started</h2>

    <p>
        1️⃣ Upload meeting transcripts or notes.
    </p>

    <p>
        2️⃣ AI automatically extracts summaries,
        action items, risks, and decisions.
    </p>

    <p>
        3️⃣ Review analytics and collaborate with your team efficiently.
    </p>

</div>
""", unsafe_allow_html=True)